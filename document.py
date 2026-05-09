import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from dotenv import load_dotenv

load_dotenv()

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'output')
ANTIEGG_CEO    = os.getenv('ANTIEGG_CEO', 'ANTIEGG 대표')
ANTIEGG_BIZ_NO = os.getenv('ANTIEGG_BIZ_NO', '000-00-00000')
ANTIEGG_PHONE  = os.getenv('ANTIEGG_PHONE', '')
ANTIEGG_EMAIL  = os.getenv('ANTIEGG_EMAIL', os.getenv('DIRECTOR_EMAIL', ''))
ANTIEGG_ADDR   = os.getenv('ANTIEGG_ADDR', '')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── 스타일 헬퍼 ────────────────────────────────────────────────────────────────

def _set_font(run, size=11, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def _cell_text(cell, text, size=10, bold=False, align='left', color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    _set_font(run, size=size, bold=bold, color=color)

def _shade_cell(cell, hex_color='1F3864'):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

# ── 금액 계산 ─────────────────────────────────────────────────────────────────

def _parse_number(value: str) -> int:
    if not value:
        return 0
    cleaned = ''.join(c for c in str(value) if c.isdigit())
    return int(cleaned) if cleaned else 0

def _fmt_money(n: int) -> str:
    return f'{n:,}원'

# ── 견적서 생성 ───────────────────────────────────────────────────────────────

def _next_quote_slot(deal: dict) -> tuple:
    """(컬럼명, 파일경로) 반환. v3 초과 시 ValueError."""
    for v in (1, 2, 3):
        col = f'quote_path_v{v}'
        if not deal.get(col):
            fname = f'{deal["deal_id"]}_quote_v{v}.docx'
            return col, os.path.join(OUTPUT_DIR, fname)
    raise ValueError('견적서 버전 초과 (v3까지만 지원)')

def generate_quote(deal: dict) -> tuple:
    """견적서 생성. (저장 컬럼명, 파일 경로) 반환."""
    col, path = _next_quote_slot(deal)

    unit_price  = _parse_number(deal.get('cond_unit_price'))
    quantity    = _parse_number(deal.get('cond_quantity')) or 1
    subtotal    = unit_price * quantity
    vat         = int(subtotal * 0.1)
    total       = subtotal + vat

    today       = datetime.now()
    valid_until = today + timedelta(days=30)
    quote_date  = today.strftime('%Y년 %m월 %d일')
    valid_str   = valid_until.strftime('%Y년 %m월 %d일')

    doc = Document()

    # ── 여백 설정 ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 제목 ──────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('견  적  서')
    _set_font(title_run, size=22, bold=True, color=(31, 56, 100))
    doc.add_paragraph()

    # ── 수신 / 견적 정보 테이블 (2열) ─────────────────────────────────────────
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    _set_col_widths(info_table, [2.5, 5.5, 2.5, 5.5])

    def info_row(r, label1, val1, label2, val2):
        cells = info_table.rows[r].cells
        _cell_text(cells[0], label1, bold=True, align='center', color=(255,255,255))
        _shade_cell(cells[0])
        _cell_text(cells[1], val1)
        _cell_text(cells[2], label2, bold=True, align='center', color=(255,255,255))
        _shade_cell(cells[2])
        _cell_text(cells[3], val2)

    info_row(0, '수 신', deal.get('company', ''), '견적번호', deal['deal_id'])
    info_row(1, '담당자', deal.get('contact_name', ''), '견적일', quote_date)
    info_row(2, '이메일', deal.get('email', ''), '유효기간', valid_str)
    info_row(3, '연락처', deal.get('contact_phone', ''), '발신', ANTIEGG_CEO)

    doc.add_paragraph()

    # ── 서비스 명세 테이블 ────────────────────────────────────────────────────
    doc.add_heading('서비스 명세', level=2)
    item_table = doc.add_table(rows=2, cols=5)
    item_table.style = 'Table Grid'
    _set_col_widths(item_table, [5.0, 5.0, 2.5, 2.0, 3.0])

    headers = ['서비스명', '서비스 설명', '단가', '수량', '공급가액']
    for i, h in enumerate(headers):
        _cell_text(item_table.rows[0].cells[i], h, bold=True, align='center', color=(255,255,255))
        _shade_cell(item_table.rows[0].cells[i])

    row1 = item_table.rows[1].cells
    _cell_text(row1[0], deal.get('cond_service_name', ''))
    _cell_text(row1[1], deal.get('cond_service_desc', ''))
    _cell_text(row1[2], _fmt_money(unit_price), align='right')
    _cell_text(row1[3], str(quantity), align='center')
    _cell_text(row1[4], _fmt_money(subtotal), align='right')

    doc.add_paragraph()

    # ── 합계 테이블 ───────────────────────────────────────────────────────────
    total_table = doc.add_table(rows=3, cols=2)
    total_table.style = 'Table Grid'
    _set_col_widths(total_table, [13.5, 4.0])

    total_rows = [
        ('공급가액', _fmt_money(subtotal)),
        ('부가세 (10%)', _fmt_money(vat)),
        ('합  계', _fmt_money(total)),
    ]
    for i, (label, value) in enumerate(total_rows):
        cells = total_table.rows[i].cells
        is_total = (i == 2)
        _cell_text(cells[0], label, bold=is_total, align='right', color=(255,255,255) if is_total else None)
        _cell_text(cells[1], value, bold=is_total, align='right', color=(255,255,255) if is_total else None)
        if is_total:
            _shade_cell(cells[0])
            _shade_cell(cells[1])

    doc.add_paragraph()

    # ── 계약 조건 ─────────────────────────────────────────────────────────────
    conditions = [
        ('결제 조건', deal.get('cond_payment_terms', '')),
        ('납품 범위', deal.get('cond_delivery_scope', '')),
        ('특이사항', deal.get('cond_notes', '')),
    ]
    has_conditions = any(v for _, v in conditions)
    if has_conditions:
        doc.add_heading('계약 조건', level=2)
        cond_table = doc.add_table(rows=len(conditions), cols=2)
        cond_table.style = 'Table Grid'
        _set_col_widths(cond_table, [3.5, 14.0])
        for i, (label, value) in enumerate(conditions):
            cells = cond_table.rows[i].cells
            _cell_text(cells[0], label, bold=True, align='center', color=(255,255,255))
            _shade_cell(cells[0], '2F5496')
            _cell_text(cells[1], value or '—')
        doc.add_paragraph()

    # ── 발신 정보 ─────────────────────────────────────────────────────────────
    sender_para = doc.add_paragraph()
    sender_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sender_lines = ['ANTIEGG']
    if ANTIEGG_ADDR:   sender_lines.append(ANTIEGG_ADDR)
    if ANTIEGG_BIZ_NO: sender_lines.append(f'사업자등록번호: {ANTIEGG_BIZ_NO}')
    if ANTIEGG_PHONE:  sender_lines.append(f'Tel: {ANTIEGG_PHONE}')
    if ANTIEGG_EMAIL:  sender_lines.append(f'Email: {ANTIEGG_EMAIL}')
    sender_run = sender_para.add_run('\n'.join(sender_lines))
    _set_font(sender_run, size=9, color=(89, 89, 89))

    doc.save(path)
    return col, path

# ── 계약서 생성 ───────────────────────────────────────────────────────────────

def _next_contract_slot(deal: dict) -> tuple:
    """(컬럼명, 파일경로) 반환. v3 초과 시 ValueError."""
    for v in (1, 2, 3):
        col = f'contract_path_v{v}'
        if not deal.get(col):
            fname = f'{deal["deal_id"]}_contract_v{v}.docx'
            return col, os.path.join(OUTPUT_DIR, fname)
    raise ValueError('계약서 버전 초과 (v3까지만 지원)')

def _article(doc, title: str, body: str):
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    _set_font(title_run, size=11, bold=True, color=(31, 56, 100))
    body_para = doc.add_paragraph()
    body_run = body_para.add_run(body)
    _set_font(body_run, size=10)

def generate_contract(deal: dict) -> tuple:
    """계약서 생성. (저장 컬럼명, 파일 경로) 반환."""
    col, path = _next_contract_slot(deal)

    unit_price = _parse_number(deal.get('cond_unit_price'))
    quantity   = _parse_number(deal.get('cond_quantity')) or 1
    subtotal   = unit_price * quantity
    vat        = int(subtotal * 0.1)
    total      = subtotal + vat

    today_str      = datetime.now().strftime('%Y년 %m월 %d일')
    contract_start = deal.get('cond_contract_start') or '___년 ___월 ___일'
    contract_end   = deal.get('cond_contract_end')   or '___년 ___월 ___일'
    client_company = deal.get('company', '')
    client_ceo     = deal.get('cond_company_ceo', '')
    client_biz_no  = deal.get('cond_company_biz_no', '')

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ── 제목 ──────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('용  역  계  약  서')
    _set_font(title_run, size=20, bold=True, color=(31, 56, 100))
    doc.add_paragraph()

    # ── 계약 당사자 테이블 ────────────────────────────────────────────────────
    party_table = doc.add_table(rows=5, cols=4)
    party_table.style = 'Table Grid'
    _set_col_widths(party_table, [1.8, 5.2, 1.8, 5.2])

    hrow = party_table.rows[0]
    hrow.cells[0].merge(hrow.cells[1])
    hrow.cells[2].merge(hrow.cells[3])
    _cell_text(hrow.cells[0], '갑 (발주자)', bold=True, align='center', color=(255, 255, 255))
    _shade_cell(hrow.cells[0])
    _cell_text(hrow.cells[2], '을 (수급자)', bold=True, align='center', color=(255, 255, 255))
    _shade_cell(hrow.cells[2])

    party_data = [
        ('상    호', client_company,                       '상    호', 'ANTIEGG'),
        ('대 표 자', client_ceo,                           '대 표 자', ANTIEGG_CEO),
        ('사업자번호', client_biz_no,                     '사업자번호', ANTIEGG_BIZ_NO),
        ('연  락  처', deal.get('contact_phone', '') or deal.get('email', ''),
                                                           '연  락  처', ANTIEGG_PHONE or ANTIEGG_EMAIL),
    ]
    for i, (la, va, lb, vb) in enumerate(party_data, start=1):
        cells = party_table.rows[i].cells
        _cell_text(cells[0], la, bold=True, align='center')
        _cell_text(cells[1], va)
        _cell_text(cells[2], lb, bold=True, align='center')
        _cell_text(cells[3], vb)

    doc.add_paragraph()

    # ── 계약 조항 ─────────────────────────────────────────────────────────────
    _article(doc, '제1조 (목적)',
        f'본 계약은 {client_company}(이하 "갑"이라 한다)와 ANTIEGG(이하 "을"이라 한다) 간에 '
        f'{deal.get("cond_service_name", "서비스")} 제공에 관한 권리와 의무를 규정함을 목적으로 한다.'
    )

    scope_text = (
        f'1. 서비스명: {deal.get("cond_service_name", "")}\n'
        f'2. 내용: {deal.get("cond_service_desc", "")}\n'
        f'3. 납품 범위: {deal.get("cond_delivery_scope", "") or "별도 협의"}'
    )
    _article(doc, '제2조 (용역 범위)', scope_text)

    _article(doc, '제3조 (계약 기간)',
        f'계약 기간은 {contract_start}부터 {contract_end}까지로 한다.'
    )

    _article(doc, '제4조 (계약 금액)',
        f'공급가액: {_fmt_money(subtotal)}\n'
        f'부가가치세 (10%): {_fmt_money(vat)}\n'
        f'합계 (VAT 포함): {_fmt_money(total)}'
    )

    _article(doc, '제5조 (결제 조건)',
        deal.get('cond_payment_terms', '계약 체결 후 별도 협의')
    )

    article_no = 6
    if deal.get('cond_notes'):
        _article(doc, f'제{article_no}조 (특이사항)', deal['cond_notes'])
        article_no += 1

    _article(doc, f'제{article_no}조 (일반 조항)',
        '① 본 계약에 명시되지 않은 사항은 관련 법령 및 상관례에 따른다.\n'
        '② 본 계약과 관련하여 분쟁이 발생한 경우 갑, 을이 협의하여 해결한다.'
    )

    doc.add_paragraph()

    # ── 체결일 ────────────────────────────────────────────────────────────────
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(today_str)
    _set_font(date_run, size=11, bold=True)

    doc.add_paragraph()

    # ── 서명란 ────────────────────────────────────────────────────────────────
    sig_table = doc.add_table(rows=4, cols=4)
    sig_table.style = 'Table Grid'
    _set_col_widths(sig_table, [1.8, 5.2, 1.8, 5.2])

    sh = sig_table.rows[0]
    sh.cells[0].merge(sh.cells[1])
    sh.cells[2].merge(sh.cells[3])
    _cell_text(sh.cells[0], '갑 (발주자)', bold=True, align='center', color=(255, 255, 255))
    _shade_cell(sh.cells[0])
    _cell_text(sh.cells[2], '을 (수급자)', bold=True, align='center', color=(255, 255, 255))
    _shade_cell(sh.cells[2])

    sig_data = [
        ('상    호', client_company,                '상    호', 'ANTIEGG'),
        ('대 표 자', f'{client_ceo}  (인)',         '대 표 자', f'{ANTIEGG_CEO}  (인)'),
        ('사업자번호', client_biz_no,              '사업자번호', ANTIEGG_BIZ_NO),
    ]
    for i, (la, va, lb, vb) in enumerate(sig_data, start=1):
        cells = sig_table.rows[i].cells
        _cell_text(cells[0], la, bold=True, align='center')
        _cell_text(cells[1], va)
        _cell_text(cells[2], lb, bold=True, align='center')
        _cell_text(cells[3], vb)

    doc.save(path)
    return col, path
